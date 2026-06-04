from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Пример отчета по преддипломной практике.docx"
OUTPUT = ROOT / "Отчет по преддипломной практике - Воронков И.С..docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def clear_document(document):
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_document_defaults(document):
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(0)


def p(document, text="", *, align=None, bold=False, size=14, indent=True):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return paragraph


def heading(document, text, level=1):
    size = 14
    paragraph = p(document, text, align=WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=size, indent=False)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def bullet(document, text):
    paragraph = p(document, text, indent=False)
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    return paragraph


def caption(document, text):
    return p(document, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, indent=False)


def figure_placeholder(document, title):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_text(cell, "[Заглушка: {}]".format(title))
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(document, "", indent=False)


def simple_table(document, headers, rows, caption_text):
    caption(document, caption_text)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    p(document, "", indent=False)
    return table


def add_title_page(document):
    p(document, "МИНОБРНАУКИ РОССИИ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    p(document, "Федеральное государственное автономное образовательное", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "учреждение высшего образования", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "«ЮЖНЫЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ»", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    p(document, "Институт компьютерных технологий и информационной безопасности", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "Направление подготовки   09.04.04 «Программная инженерия»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "", indent=False)
    p(document, "ОТЧЕТ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, indent=False)
    p(document, "о прохождении практики", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "Тема: «Мобильное приложение для анализа безопасности Android-приложений»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "обучающегося  __2__ курса", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(document, "Фамилия ____Воронков______________________________________________________", indent=False)
    p(document, "Имя _______Иван___________________________________________________________", indent=False)
    p(document, "Отчество ______Сергеевич_________________________________________________", indent=False)
    p(document, "Место практики   ___ООО «Центр исследований и разработки»________________", indent=False)
    p(document, "наименование профильной  организации", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, indent=False)
    p(document, "Вид практики ___производственная_________________________________________", indent=False)
    p(document, "Тип  практики _производственная практика (проектно-технологическая практика) ____", indent=False)
    p(document, "Способ проведения практики _стационарная___", indent=False)
    p(document, "Форма  проведения практики __ дискретная (по видам и периодам практик)___", indent=False)
    p(document, "Сроки прохождения практики  с _09.02.26____ по __06.05.26_______", indent=False)
    p(document, "Задание обучающегося на практику согласовано:", indent=False)
    p(document, "Руководитель практики", indent=False)
    p(document, "от  Университета\t\t\t\t __________/_________________", indent=False)
    p(document, "подпись                   должность, расшифровка подписи", size=10, indent=False)
    p(document, "Руководитель практики", indent=False)
    p(document, "от  профильной организации\t\t __________/_________________", indent=False)
    p(document, "подпись                   должность, расшифровка подписи", size=10, indent=False)


def add_front_matter(document):
    heading(document, "I. ЗАДАНИЕ ОБУЧАЮЩЕГОСЯ НА ПРАКТИКУ", level=0)
    tasks = [
        "Прохождение инструктажа по технике безопасности",
        "Анализ предметной области статического анализа безопасности Android-приложений",
        "Исследование существующих решений для анализа APK-файлов",
        "Проектирование архитектуры мобильного приложения для анализа безопасности Android-приложений",
        "Разработка пользовательского интерфейса мобильного приложения",
        "Разработка модуля декомпиляции APK и извлечения smali-кода",
        "Разработка rule-based модуля поиска уязвимостей в smali-коде",
        "Разработка механизма выбора security-релевантных smali-срезов",
        "Интеграция AI-анализа для уточнения и дополнения результатов",
        "Разработка настроек AI-анализа и fallback-цепочки",
        "Разработка экрана просмотра результатов и экспорта отчетов",
        "Анализ выбранных классов уязвимостей и ограничений текущей реализации",
        "Оформление отчёта по практике",
    ]
    for item in tasks:
        bullet(document, item)

    heading(document, "II. ИНСТРУКТАЖ ПО ОЗНАКОМЛЕНИЮ С ТРЕБОВАНИЯМИ ОХРАНЫ ТРУДА, ТЕХНИКИ БЕЗОПАСНОСТИ, ПОЖАРНОЙ БЕЗОПАСНОСТИ, ПРАВИЛАМ ВНУТРЕННЕГО РАСПОРЯДКА", level=0)
    p(document, "Инструктаж проведен")
    p(document, "Ознакомлен")
    p(document, "по требованиям охраны труда")
    p(document, "___________________Елькин Д. М._")
    p(document, "(подпись и ФИО  руководителя практики от профильной организации или руководителя практики от Университета, если практика проводится в Университете)", size=10)
    p(document, "«_09_» ____февраля__ 2026г.")
    p(document, "___________________Воронков И. С._")
    p(document, "(подпись и ФИО  обучающегося)", size=10)
    p(document, "«_09_» ____февраля___ 2026г.")
    p(document, "по техники безопасности")
    p(document, "по пожарной безопасности")
    p(document, "по правилами внутреннего трудового распорядка")

    heading(document, "III. ДНЕВНИК ПРАКТИКИ", level=0)
    simple_table(
        document,
        ["Дата", "Выполненные мероприятия в соответствии с заданием на практику"],
        [
            ["09.02.2026", "Прохождение инструктажа по технике безопасности"],
            ["09.02.2026 – 24.02.2026", "Анализ предметной области статического анализа безопасности Android-приложений"],
            ["24.02.2026 – 01.03.2026", "Исследование существующих SAST-решений и инструментов анализа APK"],
            ["01.03.2026 – 08.03.2026", "Проектирование архитектуры мобильного приложения"],
            ["08.03.2026 – 14.03.2026", "Разработка пользовательского интерфейса и навигации"],
            ["14.03.2026 – 21.03.2026", "Разработка модуля декомпиляции APK и обработки smali-кода"],
            ["21.03.2026 – 01.04.2026", "Разработка rule-based анализа уязвимостей"],
            ["01.04.2026 – 11.04.2026", "Интеграция настроек AI-анализа и удаленного анализатора"],
            ["11.04.2026 – 21.04.2026", "Интеграция on-device AI через AICore и fallback-цепочки"],
            ["21.04.2026 – 28.04.2026", "Реализация начального AI discovery по smali-срезам"],
            ["28.04.2026 – 04.05.2026", "Оформление отчёта по практике"],
        ],
        "",
    )

    heading(document, "IV. АНАЛИЗ ПРОВЕДЁННОЙ РАБОТЫ В ПЕРИОД ПРОХОЖДЕНИЯ ПРАКТИКИ ОБУЧАЮЩИМСЯ", level=0)
    p(document, "Раздел заполняется обучающимся в соответствии со спецификой практики")
    simple_table(
        document,
        ["№ п/п", "Выполненные мероприятия в соответствии с заданием на практику", "Анализ проделанной работы"],
        [
            ["1", "Анализ предметной области", "Изучены задачи статического анализа Android-приложений, особенности APK, DEX и smali-представления кода (см. Приложение А)."],
            ["2", "Проектирование программного решения", "Спроектирована архитектура мобильного приложения с разделением на UI, domain и data слои, а также определен pipeline анализа APK (см. Приложение Б)."],
            ["3", "Разработка пользовательского интерфейса", "Разработаны основные экраны выбора APK, настроек AI и просмотра результатов анализа на Jetpack Compose (см. Приложение В)."],
            ["4", "Разработка модуля декомпиляции", "Реализовано извлечение DEX-файлов из APK и их дизассемблирование в smali-код с использованием baksmali/dexlib2 (см. Приложение В)."],
            ["5", "Rule-based анализ", "Реализован первичный поиск уязвимостей по правилам для hardcoded secrets, небезопасной криптографии и WebView (см. Приложение В)."],
            ["6", "AI-анализ", "Реализованы remote AI analyzer, on-device AICore analyzer и fallback-цепочка для обогащения результатов рекомендациями (см. Приложение В)."],
            ["7", "AI discovery", "Реализована начальная версия поиска новых уязвимостей по security-релевантным smali-срезам (см. Приложение В)."],
            ["8", "Анализ безопасности", "Сформирован перечень основных классов уязвимостей, выбранных для анализа на текущем этапе (см. Приложение Г)."],
            ["9", "Оформление отчёта", "Оформлен отчёт по практике."],
        ],
        "",
    )

    heading(document, "ОТЗЫВ РУКОВОДИТЕЛЯ ПРАКТИКИ от ПРОФИЛЬНОЙ ОРГАНИЗАЦИИ", level=0)
    p(document, "_______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________", indent=False)
    p(document, "Руководитель практики")
    p(document, "от профильной организации\t\t   _____________/___Елькин Д. М.______")
    p(document, "подпись                       \tФИО", size=10)

    heading(document, "ОТЗЫВ РУКОВОДИТЕЛЯ ПРАКТИКИ от УНИВЕРСИТЕТА", level=0)
    p(document, "_______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________", indent=False)
    p(document, "Оценка_______________________________________________________________")
    p(document, "зачтено/отлично/хорошо/удовлетворительно", size=10)
    p(document, "Руководитель практики")
    p(document, "от Университета\t\t\t   _____________/_____Чумичева Л. В.___________")
    p(document, "подпись                       \t  ФИО", size=10)


def add_appendix_a(document):
    document.add_page_break()
    heading(document, "ПРИЛОЖЕНИЕ А. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ", level=0)
    heading(document, "А.1 Общая характеристика задачи анализа безопасности Android-приложений")
    p(document, "Мобильные приложения Android часто обрабатывают персональные данные, учетные записи пользователей, платежную информацию, сетевые токены и другие чувствительные сведения. Ошибки в реализации хранения данных, сетевого взаимодействия, криптографии и WebView могут приводить к утечкам информации, обходу механизмов защиты и выполнению нежелательных действий от имени пользователя.")
    p(document, "Разрабатываемый проект посвящен созданию мобильного приложения для анализа безопасности Android-приложений. Пользователь выбирает APK-файл, после чего приложение выполняет статический анализ его содержимого без запуска анализируемой программы. Основной объект анализа на текущем этапе — smali-код, получаемый из DEX-файлов APK.")
    p(document, "Статический анализ удобен тем, что не требует выполнения потенциально небезопасного приложения. Такой подход позволяет выявлять характерные признаки уязвимостей в коде, конфигурации и используемых API. При этом статический анализ имеет ограничения: он не всегда учитывает реальный контекст выполнения, поэтому результаты требуют фильтрации, дедупликации и ручной проверки.")

    heading(document, "А.2 Анализ существующих подходов")
    p(document, "В области анализа безопасности мобильных приложений используются несколько подходов: статический анализ, динамический анализ, гибридный анализ и ручной аудит. Каждый подход решает свою задачу и отличается по требованиям к инфраструктуре, точности и трудоемкости.")
    simple_table(
        document,
        ["Подход", "Суть", "Преимущества", "Ограничения"],
        [
            ["Статический анализ", "Изучение APK, манифеста, ресурсов и кода без запуска", "Безопасность, воспроизводимость, высокая скорость", "Возможны ложные срабатывания"],
            ["Динамический анализ", "Наблюдение за приложением во время выполнения", "Учитывает реальное поведение", "Нужна среда запуска и сценарии"],
            ["Гибридный анализ", "Комбинация статического и динамического подходов", "Более полный охват", "Сложность реализации"],
            ["Ручной аудит", "Экспертный анализ кода и архитектуры", "Высокая точность", "Большая трудоемкость"],
        ],
        "Таблица 1 – Сравнение подходов к анализу безопасности",
    )

    heading(document, "А.2.1 Статический анализ APK")
    p(document, "Статический анализ Android-приложений обычно включает извлечение содержимого APK, анализ AndroidManifest.xml, ресурсов, DEX-файлов и строковых констант. Для анализа DEX-файлов часто используется промежуточное smali-представление, которое сохраняет структуру классов, методов и вызовов API.")
    p(document, "На текущем этапе проекта реализован анализ smali-кода. Такой выбор позволяет работать непосредственно на устройстве пользователя и не требует наличия исходного кода анализируемого приложения.")

    heading(document, "А.2.2 Инструменты и аналоги")
    p(document, "Существуют готовые инструменты анализа Android-приложений, однако многие из них ориентированы на desktop/server-среду, требуют сложной установки или не рассчитаны на локальный мобильный сценарий использования.")
    simple_table(
        document,
        ["Инструмент", "Назначение", "Преимущества", "Ограничения"],
        [
            ["MobSF", "Комплексный mobile security framework", "Широкий набор проверок", "Обычно требует серверной/desktop-среды"],
            ["JADX", "Декомпиляция APK в Java-подобный код", "Удобен для ручного анализа", "Не является полноценным SAST-решением"],
            ["QARK", "Поиск типовых Android-уязвимостей", "Ориентация на безопасность Android", "Ограничения актуальности и покрытия"],
            ["AndroBugs", "Статический анализ APK", "Простота использования", "Ограниченное развитие и гибкость"],
            ["Android Lint", "Анализ исходного проекта Android", "Интеграция с Gradle/IDE", "Требует исходный код, а не только APK"],
        ],
        "Таблица 2 – Сравнение существующих инструментов",
    )

    heading(document, "А.3 Обоснование применения AI")
    p(document, "Rule-based анализ хорошо подходит для поиска конкретных паттернов, например строк с ключевыми словами password, token или secret. Однако не все уязвимости выражаются простым совпадением строк. В некоторых случаях необходимо учитывать контекст метода, назначение вызываемого API и связь между несколькими инструкциями smali-кода.")
    p(document, "Для повышения полезности результатов в проекте используется AI-анализ. Он применяется не как замена детерминированным правилам, а как дополнительный слой: AI уточняет риск, формирует рекомендацию и в начальной реализации может находить потенциальные уязвимости в заранее отобранных smali-срезах.")
    p(document, "Важным ограничением является то, что в модель не передается весь smali tree. Сначала выполняется deterministic candidate extraction, после чего AI получает только небольшие security-релевантные фрагменты кода.")

    heading(document, "А.4 Вывод")
    p(document, "Проведенный анализ показал, что задача мобильного SAST для Android-приложений актуальна из-за распространенности APK-файлов, сложности ручного анализа smali-кода и необходимости быстро выявлять типовые ошибки безопасности. Для текущего этапа разработки выбран гибридный подход: deterministic rule-based scan используется как основа, а AI применяется для обогащения и начального обнаружения дополнительных finding-ов.")


def add_appendix_b(document):
    document.add_page_break()
    heading(document, "ПРИЛОЖЕНИЕ Б. ПРОЕКТИРОВАНИЕ", level=0)
    heading(document, "Б.1 Общая архитектура программного решения")
    p(document, "Программное решение представляет собой Android-приложение, разработанное на Kotlin с использованием Jetpack Compose. Архитектура проекта построена с разделением на UI, domain и data слои. Для внедрения зависимостей используется Koin.")
    p(document, "UI-слой отвечает за отображение экранов, обработку пользовательских событий и навигацию. Domain-слой содержит основную бизнес-логику анализа APK, AI-обогащение результатов и генерацию отчетов. Data-слой отвечает за декомпиляцию APK, работу с настройками и извлечение smali-срезов.")
    figure_placeholder(document, "Схема взаимодействия между UI, domain и data слоями")
    caption(document, "Рисунок 1 – Схема взаимодействия между слоями")

    heading(document, "Б.2 Pipeline анализа APK")
    p(document, "Основной сценарий работы начинается с выбора пользователем APK-файла. Файл передается в AnalyzerInteractor, который управляет последовательностью анализа. Если APK выбран через SAF URI, содержимое предварительно копируется во временный файл, что позволяет корректно работать с content:// URI.")
    p(document, "После получения файла выполняется декомпиляция DEX-файлов в smali-код. Далее SmaliAnalyzer выполняет rule-based scan. Если AI-анализ включен в настройках, дополнительно запускается SmaliVulnerabilityDiscoverer, который ищет security-релевантные smali-срезы и пытается обнаружить новые потенциальные уязвимости. После этого результаты дедуплицируются и передаются в VulnerabilityAIAnalyzer для получения AI-рекомендаций.")
    figure_placeholder(document, "Pipeline анализа APK: выбор файла, декомпиляция, scan, AI discovery, AI enrichment, отчет")
    caption(document, "Рисунок 2 – Pipeline анализа APK")

    heading(document, "Б.3 Основные функциональные модули системы")
    heading(document, "Б.3.1 Модуль пользовательского интерфейса")
    p(document, "Модуль пользовательского интерфейса включает MainScreen, SettingsScreen и DetailsScreen. MainScreen отвечает за выбор APK и отображение прогресса анализа. SettingsScreen содержит параметры AI-анализа, включая выбор режима ON_DEVICE или REMOTE. DetailsScreen отображает результаты, группирует finding-и по категориям и предоставляет экспорт отчета.")

    heading(document, "Б.3.2 Модуль декомпиляции APK")
    p(document, "ApkDecompiler извлекает DEX-файлы из APK и преобразует их в smali-представление с помощью baksmali/dexlib2. Результаты сохраняются во временную директорию cacheDir/smali_output. После завершения анализа выполняется cleanup.")

    heading(document, "Б.3.3 Модуль rule-based анализа")
    p(document, "SmaliAnalyzer проходит по smali-файлам и ищет известные признаки уязвимостей. На текущем этапе реализованы правила для HARDCODED_SECRET, INSECURE_CRYPTO и WEBVIEW_JS_ENABLED. Результатом работы модуля является поток объектов Vulnerability.")

    heading(document, "Б.3.4 Модуль AI discovery")
    p(document, "SmaliCandidateExtractor выделяет методы, содержащие security-релевантные сигналы: WebView bridge, TLS/network, cleartext HTTP, crypto, storage, logging, dynamic code loading и process execution. AICoreSmaliDiscoverer анализирует такие срезы через ML Kit Prompt API при наличии AICore. HeuristicSmaliDiscoverer обеспечивает fallback без обращения к модели.")

    heading(document, "Б.3.5 Модуль AI enrichment")
    p(document, "AI enrichment реализован через интерфейс VulnerabilityAIAnalyzer. RemoteAIAnalyzer отправляет чанки finding-ов в OpenAI-compatible API, AICoreAnalyzer выполняет локальный анализ через AICore/Gemini Nano, а FallbackAIAnalyzer обеспечивает последовательную деградацию до доступного варианта. Результатом является VulnerabilityWithAIInsight с risk score, severity и recommendation.")
    figure_placeholder(document, "Fallback-цепочка AI-анализа: AICore, Remote, Heuristics")
    caption(document, "Рисунок 3 – Схема fallback-цепочки AI-анализа")

    heading(document, "Б.3.6 Модуль отчетов")
    p(document, "ReportGenerator формирует отчеты в форматах TXT, CSV и PDF. На текущем этапе экспорт содержит базовые сведения о найденных уязвимостях: файл, строку, тип и описание. Расширение качества отчетов, включая полное отображение AI metadata, запланировано как дальнейшая работа.")

    heading(document, "Б.4 Модель данных")
    p(document, "Основные доменные модели: Vulnerability, AnalysisReport и VulnerabilityWithAIInsight. Vulnerability описывает конкретный finding: тип, описание, файл и строку. AnalysisReport хранит путь к APK, список уязвимостей, текстовое summary и список AI-insight-ов. VulnerabilityWithAIInsight связывает исходную уязвимость с оценкой риска, уровнем severity и рекомендацией.")

    heading(document, "Б.5 Акторы и сценарии использования")
    p(document, "Основным актором является пользователь мобильного приложения. Пользователь может выбрать APK-файл, запустить анализ, просмотреть сгруппированные результаты, открыть настройки AI, отфильтровать результаты поиска и экспортировать отчет.")
    p(document, "Дополнительным актором можно считать специалиста по информационной безопасности, который использует полученный отчет как первичный список finding-ов для дальнейшей ручной проверки и приоритизации исправлений.")


def add_appendix_c(document):
    document.add_page_break()
    heading(document, "ПРИЛОЖЕНИЕ В. РАЗРАБОТКА ПРОГРАММНОГО РЕШЕНИЯ", level=0)
    heading(document, "В.1 Разработка пользовательского интерфейса мобильного приложения")
    p(document, "Пользовательский интерфейс разработан с использованием Jetpack Compose. Такой подход позволяет описывать интерфейс декларативно и связывать его с состоянием MainUiState и настройками приложения.")

    heading(document, "В.1.1 Главный экран")
    p(document, "MainScreen предназначен для выбора APK-файла и отображения текущего состояния анализа. Во время анализа экран показывает stage и progress. Для предотвращения выхода во время загрузки используется PredictiveBackHandler.")
    figure_placeholder(document, "Скриншот главного экрана выбора APK")
    caption(document, "Рисунок 4 – Интерфейс главного экрана")

    heading(document, "В.1.2 Экран настроек AI")
    p(document, "SettingsScreen позволяет включить или отключить AI-анализ, выбрать режим ON_DEVICE или REMOTE, настроить URL модели, API key, название remote-модели и размер чанка. Настройки сохраняются через SettingsRepository на базе DataStore Preferences.")
    figure_placeholder(document, "Скриншот экрана настроек AI")
    caption(document, "Рисунок 5 – Интерфейс экрана настроек AI")

    heading(document, "В.1.3 Экран результатов анализа")
    p(document, "DetailsScreen отображает итоговый AnalysisReport. Результаты группируются по категориям, поддерживаются раскрытие категорий, поиск по файлу/описанию, отображение AI Security Analysis и экспорт отчета.")
    figure_placeholder(document, "Скриншот экрана результатов анализа")
    caption(document, "Рисунок 6 – Интерфейс экрана результатов анализа")

    heading(document, "В.2 Разработка data и domain слоёв мобильного приложения")
    heading(document, "В.2.1 Разработка data слоя")
    p(document, "Data слой включает ApkDecompiler, SmaliAnalyzer, SmaliCandidateExtractor и SettingsRepository. Эти компоненты отвечают за получение анализируемых данных, работу с временными файлами, чтение smali-кода и хранение настроек анализа.")
    p(document, "ApkDecompiler извлекает classes*.dex из APK и дизассемблирует их в smali. SmaliAnalyzer выполняет первичный rule-based анализ. SmaliCandidateExtractor выделяет методы, которые потенциально могут содержать security-релевантную логику.")

    heading(document, "В.2.2 Разработка domain слоя")
    p(document, "Domain слой содержит AnalyzerInteractor, AI-анализаторы, модели отчета и ReportGenerator. AnalyzerInteractor является центральным оркестратором pipeline: он получает APK, запускает декомпиляцию, собирает finding-и, выполняет AI discovery и AI enrichment, дедуплицирует результаты и формирует AnalysisReport.")
    p(document, "На текущем этапе AnalyzerInteractor также содержит часть логики summary, дедупликации и cleanup. В дальнейшем эту логику целесообразно вынести в отдельные pure helpers, чтобы упростить тестирование и сопровождение.")

    heading(document, "В.2.3 Взаимодействие между слоями")
    p(document, "MainViewModel подписывается на AnalyzerInteractor.analysisState и преобразует доменные состояния в состояние UI. При выборе APK ViewModel обновляет выбранное имя файла и запускает analyzeApk. После завершения анализа пользователь автоматически переходит на экран деталей.")
    figure_placeholder(document, "Диаграмма последовательности выбора APK и получения AnalysisReport")
    caption(document, "Рисунок 7 – Диаграмма последовательности анализа APK")

    heading(document, "В.3 Разработка rule-based анализа")
    p(document, "Rule-based анализ реализован в SmaliAnalyzer. Он проходит по smali-файлам и проверяет каждую строку на наличие сигнатур. Если найдено совпадение, создается объект Vulnerability с типом, описанием, путем к файлу и номером строки.")
    p(document, "Текущая реализация содержит три базовых правила: HARDCODED_SECRET, INSECURE_CRYPTO и WEBVIEW_JS_ENABLED. Эти правила выбраны как распространенные и понятные примеры проблем безопасности Android-приложений.")

    heading(document, "В.4 Разработка AI discovery")
    p(document, "AI discovery предназначен для поиска потенциальных уязвимостей, которые не покрываются простыми правилами. Для этого SmaliCandidateExtractor сначала отбирает только методы с релевантными сигналами. Затем AICoreSmaliDiscoverer отправляет каждый срез в модель и ожидает компактный JSON с типом, описанием, строкой и confidence.")
    p(document, "На текущем этапе discovery является best-effort механизмом. Результаты проходят минимальную валидацию: проверяются непустой тип, описание, confidence threshold и попадание строки в диапазон исходного smali-среза. После этого finding-и дедуплицируются с rule-based результатами.")

    heading(document, "В.5 Разработка AI enrichment")
    p(document, "AI enrichment применяется к списку уже найденных уязвимостей. AICoreAnalyzer анализирует каждую уязвимость локально через ML Kit Prompt API, если устройство поддерживает AICore. RemoteAIAnalyzer использует OpenAI-compatible endpoint и обрабатывает finding-и чанками. Если AI недоступен, используется fallback с эвристическими рекомендациями.")

    heading(document, "В.6 Разработка экспорта отчетов")
    p(document, "ReportGeneratorImpl формирует отчеты в TXT, CSV и PDF. Файлы сохраняются в Documents через MediaStore. На текущем этапе экспорт реализован в базовом виде и требует дальнейшего улучшения: более полного учета AI severity, risk score, recommendations, source и confidence.")

    heading(document, "В.7 Текущий статус реализации")
    p(document, "На текущем этапе реализованы remote AI analysis с chunking, настройки AI, отображение AI insights, fallback-цепочка AICore → Remote → Heuristics и начальная smali-based AI discovery. Следующими направлениями развития являются hardening discovery, улучшение жизненного цикла анализа, повышение качества отчетов и добавление unit-тестов на deterministic parsing/deduplication.")


def add_appendix_d(document):
    document.add_page_break()
    heading(document, "ПРИЛОЖЕНИЕ Г. БЕЗОПАСНОСТЬ", level=0)
    heading(document, "Г.1 Цель выбора классов уязвимостей")
    p(document, "Раздел посвящен обоснованию классов уязвимостей, выбранных для анализа на текущем этапе разработки мобильного SAST-приложения. Приоритет отдавался уязвимостям, которые часто встречаются в Android-приложениях, могут быть обнаружены по smali-коду и имеют практическую значимость для защиты пользовательских данных.")

    heading(document, "Г.2 Основные классы уязвимостей")
    simple_table(
        document,
        ["Класс уязвимости", "Причина выбора", "Текущий механизм анализа"],
        [
            ["HARDCODED_SECRET", "Ключи, токены и пароли в коде могут быть извлечены из APK", "Rule-based поиск const-string и чувствительных ключевых слов"],
            ["INSECURE_CRYPTO", "Слабые алгоритмы и режимы шифрования снижают стойкость защиты данных", "Rule-based поиск вызовов Cipher и небезопасных алгоритмов"],
            ["WEBVIEW_JS_ENABLED", "Небезопасная настройка WebView может привести к XSS и доступу к JavaScript bridge", "Rule-based и candidate extraction по WebView API"],
            ["INSECURE_TLS_VALIDATION", "Отключение проверки сертификатов позволяет MITM-атаки", "Heuristic discovery по TLS/HostnameVerifier/X509TrustManager сигналам"],
            ["CLEARTEXT_NETWORK", "HTTP-трафик может быть перехвачен и изменен", "Candidate extraction по http:// и HttpURLConnection"],
            ["INSECURE_STORAGE", "Чувствительные данные в незащищенном хранилище могут быть доступны третьим лицам", "Candidate extraction по SharedPreferences, files, SQLiteDatabase и sensitive words"],
            ["SENSITIVE_LOGGING", "Логи могут раскрывать токены, пароли и персональные данные", "Heuristic discovery по Android Log/System output и sensitive words"],
            ["DYNAMIC_CODE_LOADING", "Загрузка кода во время выполнения повышает риск исполнения непроверенного кода", "Candidate extraction по DexClassLoader, PathClassLoader и reflection"],
            ["PROCESS_EXECUTION", "Запуск системных команд может привести к RCE-подобным сценариям при плохой валидации ввода", "Candidate extraction по Runtime.exec и ProcessBuilder"],
        ],
        "Таблица 3 – Основные классы уязвимостей текущего этапа",
    )

    heading(document, "Г.3 Обоснование rule-based правил")
    p(document, "Rule-based правила выбраны как базовый детерминированный слой анализа. Они дают воспроизводимый результат, не зависят от доступности модели и позволяют быстро находить очевидные проблемы. HARDCODED_SECRET важен из-за простоты извлечения строк из APK. INSECURE_CRYPTO важен из-за частого использования устаревших алгоритмов. WEBVIEW_JS_ENABLED выбран из-за распространенности WebView в Android-приложениях и риска небезопасного взаимодействия JavaScript с нативным кодом.")

    heading(document, "Г.4 Обоснование AI-assisted discovery")
    p(document, "Некоторые проблемы безопасности требуют анализа контекста метода, а не только одной строки. Например, TLS bypass может быть связан с реализацией HostnameVerifier или X509TrustManager, а dynamic code loading требует понимания источника загружаемого кода. Для таких случаев используется AI-assisted discovery: сначала deterministic extractor находит релевантный smali-срез, затем AI или heuristic fallback формирует finding.")
    p(document, "Этот подход снижает риск отправки лишнего объема данных в модель и ограничивает анализ только небольшими фрагментами, содержащими security signals. На текущем этапе discovery считается экспериментальным механизмом и требует дальнейшей настройки candidate ranking, confidence metadata и тестовых fixture-ов.")

    heading(document, "Г.5 Ограничения текущей реализации")
    p(document, "Текущая версия анализирует преимущественно smali-код и не содержит полноценного manifest scanner. Поэтому такие проверки, как exported components, debuggable, cleartextTrafficPermitted и backup flags, запланированы как дальнейшее расширение. Также AI-discovered findings пока кодируют source в description через префикс AI-discovered, а structured metadata для source, confidence и evidence еще не выделены в отдельную модель.")
    p(document, "SmaliCandidateExtractor и SmaliAnalyzer на текущем этапе читают файлы через readLines(), что удобно для начальной реализации, но может быть неоптимально для больших APK. В дальнейшем целесообразно перейти к streaming/sequence чтению и добавить ограничения по количеству кандидатов с учетом severity и размера метода.")

    heading(document, "Г.6 Вывод")
    p(document, "Выбранные классы уязвимостей покрывают наиболее практичные направления первичного анализа Android-приложений: секреты, криптографию, WebView, сеть, хранилище, логи, динамическую загрузку кода и выполнение команд. На текущем этапе проект реализует базовый rule-based слой и начальную AI-assisted discovery, что формирует основу для дальнейшего расширения SAST-возможностей и повышения качества отчетов.")


def build_document():
    document = Document(TEMPLATE)
    clear_document(document)
    set_document_defaults(document)
    add_title_page(document)
    add_front_matter(document)
    add_appendix_a(document)
    add_appendix_b(document)
    add_appendix_c(document)
    add_appendix_d(document)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
